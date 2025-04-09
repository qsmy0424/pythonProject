from docx import Document

# 打开 .docx 文件
doc = Document("C:\\Users\\qsmy\\Desktop\\文书评分(1).docx")  # 替换为你的文件路径

# 遍历文档中的所有段落
for i, paragraph in enumerate(doc.paragraphs):
    # 如果段落内容不为空，则输出
    if paragraph.text.strip():
        print(f"段落 {i + 1}: {paragraph.text}")

    # 如果当前段落下面为表格，则输出
    # if i < len(doc.paragraphs) - 1 and doc.paragraphs[i + 1].text.strip():
    #     print(f"段落 {i + 2} 为表格，输出表格内容：")
    #     for row in doc.tables[i].rows:
    #         row_data = [cell.text.strip() for cell in row.cells]
    #         print(row_data)
    #     print()
