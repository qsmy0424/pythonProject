import sys
import docx
import openpyxl
from openpyxl.styles import Alignment
from openpyxl.styles import Font

# if len(sys.argv) < 3:
#     exit(0)

# file = docx.Document(sys.argv[1])
file = docx.Document("E:\\1234.docx")

# print(len(file.tables))

# typedict = {}


# for index, table in enumerate(file.tables):
#     print("Table {}: {}".format(index, table))
#     for i in range(len(table.rows)):
#         for j in range(len(table.columns)):
#             # print(f"{i}行{j}列：数据：{table.cell(i, j).text}")
#             if table.cell(i, j).text == '题目内容' and '题目内容' not in typedict:
#                 typedict['题目内容'] = j
#             if table.cell(i, j).text == '题目选项' and '题目选项' not in typedict:
#                 typedict['题目选项'] = j
#             if table.cell(i, j).text == '题目答案' and '题目答案' not in typedict:
#                 typedict['题目答案'] = j
#             if table.cell(i, j).text == '难度系数' and '难度系数' not in typedict:
#                 typedict['难度系数'] = j

# print(typedict)
# print(typedict.values())

# for index, table in enumerate(file.tables):
#     print("Table {}: {}".format(index, table))
#     for row in table.rows:
#         if index == 0:
#             print(row.cells[0].text)
#             print(row.cells[1].text)
#             print(row.cells[3].text)
#             print(row.cells[5].text)
#         if index == 1:
#             print(row.cells[0].text)
#             print(row.cells[2].text)
#             print(row.cells[4].text)
#             print(row.cells[6].text)
#         if index == 2:
#             print(row.cells[0].text)
#             print(row.cells[2].text)
#             print(row.cells[4].text)
#         print("    ")

QUESTION_CONTENT = '题目内容'
QUESTION_OPTION = '题目选项'
QUESTION_ANSWER = '题目答案'
QUESTION_DIFFICULTY = '难度系数'

# df = openpyxl.Workbook()
# df = openpyxl.load_workbook(r'E:\\export.xlsx')
# df = openpyxl.load_workbook(sys.argv[1])
df = openpyxl.load_workbook("E:\\export.xlsx")
sheet = df.active

for index, table in enumerate(file.tables):
    if index == 0:
        question_type = '单选文本题'
    elif index == 1:
        question_type = '多选文本题'
    else:
        question_type = '是非题'

    column_dict = {QUESTION_CONTENT: 0, QUESTION_OPTION: 1, QUESTION_ANSWER: 2, QUESTION_DIFFICULTY: 3}

    # for row in table.rows:
    #     if QUESTION_CONTENT not in row.cells[0].text:
    #         continue
    #     else:
    #         for column_index, cell in enumerate(row.cells):
    #             if cell.text not in column_dict:
    #                 column_dict[cell.text] = column_index
    #         break
    for row in table.rows:
        question = row.cells[column_dict.get(QUESTION_CONTENT)].text
        if len(question) == 0:
            continue
        if question[0].isdigit():
            new_list = []
            if '、' in question:
                question = question.replace('、、', '')
                question = question[question.find('、') + 1:]
            else:
                if question[0].isdigit():
                    question = question[1:]
                    if question[0].isdigit():
                        question = question[1:]
            if index == 0 or index == 1:
                option_list = row.cells[column_dict.get(QUESTION_OPTION)].text.split('\n')

                new_list = [question_type, row.cells[column_dict.get(QUESTION_DIFFICULTY)].text.strip(), question]

                for option in option_list:
                    if len(option) != 0:
                        new_list.append(option[2:].strip())

                new_list.append("")
                new_list.append("")
                if index == 0:
                    new_list.append(row.cells[column_dict.get(QUESTION_ANSWER)].text.strip())
                else:
                    new_list.append(
                        "|".join([char for char in row.cells[column_dict.get(QUESTION_ANSWER)].text.strip()]))

            if index == 2:
                new_list = [question_type, row.cells[2].text.strip(), question, "",
                            "", "", "", "", "", row.cells[1].text.strip()]
            sheet.append(new_list)

font = sheet['D1'].font
for index, rows in enumerate(sheet.rows):
    if index > 0:
        for column_index, cell in enumerate(rows):
            cell.font = Font(name=font.name, size=font.size, bold=font.bold, italic=font.italic, color=font.color)
            if column_index > 0:
                cell.alignment = Alignment(wrapText=True)
            else:
                cell.alignment = Alignment(horizontal='center', vertical='center')

df.save('result.xlsx')
