import sys
import docx
import re
import openpyxl
from openpyxl.styles import Alignment
from openpyxl.styles import Font
from openpyxl import Workbook

# 导入试题的文件
file = docx.Document("E:\GZ062 法律实务赛题及答案（模块一）(1).docx")

# 答案选项，如果ABCD选项后面跟的是顿号，则改为'[A-D]、'
pattern = '[A-D].'

workbook = Workbook()
sheet = workbook.active

option_list = []

# for index, paragraph in enumerate(file.paragraphs):
row_index = 1
for paragraph in file.paragraphs:
    if paragraph.text:
        if paragraph.text[0].isdigit():
            sheet.cell(row=row_index, column=1).value = paragraph.text
            #print(paragraph.text)
            print("")
        # 如果答案选项后面为顿号，则下面的也要改
        elif (paragraph.text.startswith('A.')
              or paragraph.text.startswith('B.')
              or paragraph.text.startswith('C.')
              or paragraph.text.startswith('D.')):
            if len(option_list) == 4:
                for index, option in enumerate(option_list):
                    sheet.cell(row=row_index, column=4 + index).value = option

                option_list.clear()
                row_index += 1
            temp_list = re.split(pattern, paragraph.text)
            temp_list = [char.strip() for char in temp_list if len(char) != 0]
            option_list.extend(temp_list)


workbook.save("export.xlsx")
