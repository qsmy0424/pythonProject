import sys
import docx
import re
import openpyxl
from openpyxl.styles import Alignment
from openpyxl.styles import Font

if len(sys.argv) < 3:
    exit(0)

file = docx.Document(sys.argv[1])
# file = docx.Document("E:\\12345.docx")


QUESTION_CONTENT = '题目内容'
QUESTION_OPTION = '题目选项'
QUESTION_ANSWER = '题目答案'
QUESTION_DIFFICULTY = '难度系数'

# 构建正则表达式
pattern = '[A-D]\.'

# df = openpyxl.Workbook()
# df = openpyxl.load_workbook(r'E:\\export.xlsx')
df = openpyxl.load_workbook(sys.argv[2])
# df = openpyxl.load_workbook("E:\\export.xlsx")
sheet = df.active

question_type = ''
for index, table in enumerate(file.tables):

    column_dict = {QUESTION_CONTENT: 0, QUESTION_OPTION: 1, QUESTION_ANSWER: 2, QUESTION_DIFFICULTY: 3}

    # for row in table.rows:
    #     if QUESTION_CONTENT not in row.cells[0].text:
    #         continue
    #     else:
    #         for column_index, cell in enumerate(row.cells):
    #             if cell.text not in column_dict:
    #                 column_dict[cell.text] = column_index
    #         break
    for row in table.rows:
        question = row.cells[column_dict.get(QUESTION_CONTENT)].text
        if len(question) == 0:
            continue
        if question[0].isdigit():
            new_list = []
            # if '、' in question:
            #     question = question.replace('、、', '')
            #     question = question[question.find('、') + 1:]
            # else:
            #     if question[0].isdigit():
            #         question = question[1:]
            #         if question[0].isdigit():
            #             question = question[1:]
            if question[0].isdigit():
                question = question[1:]
                if question[0].isdigit():
                    question = question[1:]
                if question[0] == '、' or question[0] == '.':
                    question = question[1:]

        if len(row.cells) == 4:
            # option_list = row.cells[column_dict.get(QUESTION_OPTION)].text.split('\n')
            option_list = re.split(pattern, row.cells[column_dict.get(QUESTION_OPTION)].text)
            answer = row.cells[column_dict.get(QUESTION_ANSWER)].text.strip()

            if len(answer) == 1:
                question_type = '单选文本题'
            else:
                question_type = '多选文本题'

            new_list = [question_type, row.cells[column_dict.get(QUESTION_DIFFICULTY)].text.strip(), question]

            for option in option_list:
                option = option.strip()
                if len(option) != 0:
                    new_list.append(option)

            new_list.append("")
            new_list.append("")
            if len(answer) == 1:
                new_list.append(answer)
            else:
                new_list.append(
                    "|".join([char for char in answer]))
        else:
            question_type = '是非题'
            new_list = [question_type, row.cells[2].text.strip(), question, "",
                        "", "", "", "", "", row.cells[1].text.strip()]
        sheet.append(new_list)

font = sheet['D1'].font
for index, rows in enumerate(sheet.rows):
    if index > 0:
        for column_index, cell in enumerate(rows):
            cell.font = Font(name=font.name, size=font.size, bold=font.bold, italic=font.italic, color=font.color)
            if column_index > 0:
                cell.alignment = Alignment(wrapText=True)
            else:
                cell.alignment = Alignment(horizontal='center', vertical='center')

df.save('result.xlsx')
